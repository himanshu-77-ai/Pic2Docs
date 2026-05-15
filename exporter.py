"""
exporter.py — Production Export Pipeline
─────────────────────────────────────────
All exports return raw bytes (via BytesIO) — NO disk writes.
Handles:  TXT · PDF · DOCX · XLSX · Notebook-style PNG
Each function returns (bytes | None, error_message | None).
"""

from __future__ import annotations

import io
import logging
import textwrap
from datetime import datetime
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from PIL import Image, ImageDraw, ImageFont

logger = logging.getLogger("pic2docs.exporter")

# Path to the bundled Unicode font (DejaVuSans.ttf must be in project root)
_FONT_PATH = Path(__file__).parent / "DejaVuSans.ttf"

ExportResult = tuple[Optional[bytes], Optional[str]]   # (data, error)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _timestamp() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _font_path_str() -> str:
    """Return font path as string; log warning if missing."""
    if not _FONT_PATH.exists():
        logger.warning("DejaVuSans.ttf not found at %s — PDF may use fallback font.", _FONT_PATH)
    return str(_FONT_PATH)


# ── TXT export ────────────────────────────────────────────────────────────────

def export_txt(text: str, filename: str = "extracted_text") -> ExportResult:
    """Plain UTF-8 text file with header metadata."""
    try:
        header = (
            f"Pic2Docs — Extracted Text\n"
            f"Generated: {_timestamp()}\n"
            f"Source file: {filename}\n"
            f"{'─' * 60}\n\n"
        )
        content = (header + text).encode("utf-8")
        return content, None
    except Exception as exc:
        logger.exception("TXT export failed: %s", exc)
        return None, f"TXT export failed: {exc}"


# ── PDF export ────────────────────────────────────────────────────────────────

class _Pic2DocsPDF(FPDF):
    """Custom FPDF subclass with header and footer."""

    def header(self):
        self.set_font("DejaVu", size=9)
        self.set_text_color(120, 120, 140)
        self.cell(0, 8, "Pic2Docs — Extracted Text", align="L")
        self.ln(4)
        self.set_draw_color(200, 200, 220)
        self.set_line_width(0.3)
        self.line(self.l_margin, self.get_y(), 210 - self.r_margin, self.get_y())
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        self.set_font("DejaVu", size=8)
        self.set_text_color(150, 150, 160)
        self.cell(0, 10, f"Page {self.page_no()} | Generated {_timestamp()}", align="C")


def export_pdf(text: str, filename: str = "extracted_text") -> ExportResult:
    """
    Professional notebook-style PDF.
    Uses DejaVuSans for full Unicode support (Latin, Devanagari, Bengali, CJK, Arabic).
    No disk writes — returns bytes directly.
    """
    try:
        pdf = _Pic2DocsPDF()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.set_margins(left=20, top=20, right=20)

        # Register the bundled Unicode font
        font_path = _font_path_str()
        pdf.add_font("DejaVu", style="", fname=font_path, uni=True)

        pdf.add_page()
        pdf.set_font("DejaVu", size=12)
        pdf.set_text_color(30, 30, 40)

        # Title block
        pdf.set_font("DejaVu", size=18)
        pdf.set_text_color(80, 70, 200)
        pdf.cell(0, 12, "Extracted Text", ln=True)
        pdf.set_font("DejaVu", size=9)
        pdf.set_text_color(120, 120, 140)
        pdf.cell(0, 6, f"Source: {filename}  |  {_timestamp()}", ln=True)
        pdf.ln(4)
        pdf.set_draw_color(180, 170, 240)
        pdf.set_line_width(0.5)
        pdf.line(pdf.l_margin, pdf.get_y(), 210 - pdf.r_margin, pdf.get_y())
        pdf.ln(8)

        # Body
        pdf.set_font("DejaVu", size=12)
        pdf.set_text_color(30, 30, 40)
        for line in text.split("\n"):
            if line.strip() == "":
                pdf.ln(4)
            else:
                pdf.multi_cell(0, 7, txt=line, align="L")

        # Return bytes directly (fpdf2 >= 2.7 returns bytes from output())
        raw = pdf.output()
        if isinstance(raw, str):
            raw = raw.encode("latin-1")
        return bytes(raw), None

    except Exception as exc:
        logger.exception("PDF export failed: %s", exc)
        return None, f"PDF export failed: {exc}"


# ── DOCX export ───────────────────────────────────────────────────────────────

def export_docx(text: str, filename: str = "extracted_text") -> ExportResult:
    """
    Professional Word document with metadata, styled heading, and editable body.
    Fully in-memory — no disk writes.
    """
    try:
        doc = Document()

        # Document properties
        core = doc.core_properties
        core.title = "Pic2Docs — Extracted Text"
        core.subject = f"OCR output from {filename}"
        core.author = "Pic2Docs"
        core.comments = f"Generated at {_timestamp()}"

        # Custom heading style
        h = doc.add_heading("Extracted Text", level=1)
        h.runs[0].font.color.rgb = RGBColor(80, 70, 200)
        h.runs[0].font.size = Pt(22)

        # Metadata line
        meta = doc.add_paragraph()
        meta.paragraph_format.space_after = Pt(12)
        run = meta.add_run(f"Source: {filename}  |  Generated: {_timestamp()}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 140)
        run.italic = True

        # Divider (horizontal rule via border)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        p_rule = doc.add_paragraph()
        pPr = p_rule._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "B0A8E8")
        pBdr.append(bottom)
        pPr.append(pBdr)

        # Body text — each OCR line as its own paragraph
        for line in text.split("\n"):
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"

        # Set default font for the document body
        from docx.oxml.ns import qn as _qn
        styles = doc.styles["Normal"]
        styles.font.name = "Calibri"
        styles.font.size = Pt(11)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue(), None

    except Exception as exc:
        logger.exception("DOCX export failed: %s", exc)
        return None, f"DOCX export failed: {exc}"


# ── XLSX export ───────────────────────────────────────────────────────────────

def export_xlsx(text: str, filename: str = "extracted_text") -> ExportResult:
    """
    Formatted Excel workbook with:
      - Sheet 1: Line-by-line extracted text with styled header
      - Sheet 2: Word frequency analysis
    Fully in-memory — no disk writes.
    """
    try:
        wb = Workbook()

        # ── Sheet 1: Extracted Text ──
        ws1 = wb.active
        ws1.title = "Extracted Text"

        # Header style
        header_fill   = PatternFill("solid", fgColor="5046E4")
        header_font   = Font(name="Calibri", bold=True, color="FFFFFF", size=12)
        subhead_fill  = PatternFill("solid", fgColor="EEEDff")
        subhead_font  = Font(name="Calibri", italic=True, color="5046E4", size=9)
        body_font     = Font(name="Calibri", size=11)
        border_side   = Side(style="thin", color="D0CEEE")
        body_border   = Border(
            left=border_side, right=border_side,
            top=border_side, bottom=border_side
        )

        # Title row
        ws1.merge_cells("A1:C1")
        title_cell = ws1["A1"]
        title_cell.value = "Pic2Docs — Extracted Text"
        title_cell.font = header_font
        title_cell.fill = header_fill
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws1.row_dimensions[1].height = 24

        # Meta row
        ws1.merge_cells("A2:C2")
        meta_cell = ws1["A2"]
        meta_cell.value = f"Source: {filename}  |  Generated: {_timestamp()}"
        meta_cell.font = subhead_font
        meta_cell.fill = subhead_fill
        meta_cell.alignment = Alignment(horizontal="left", vertical="center")
        ws1.row_dimensions[2].height = 18

        # Column headers
        col_headers = ["#", "Text Line", "Word Count"]
        for col_idx, header in enumerate(col_headers, start=1):
            cell = ws1.cell(row=3, column=col_idx, value=header)
            cell.font = Font(name="Calibri", bold=True, color="5046E4", size=10)
            cell.fill = PatternFill("solid", fgColor="F0EFFF")
            cell.alignment = Alignment(horizontal="center")
            cell.border = body_border

        # Data rows
        lines = [ln for ln in text.split("\n")]
        alt_fill = PatternFill("solid", fgColor="F8F8FF")
        for row_idx, line in enumerate(lines, start=4):
            word_count = len(line.split()) if line.strip() else 0
            row_fill = alt_fill if row_idx % 2 == 0 else None
            for col_idx, value in enumerate([row_idx - 3, line, word_count], start=1):
                cell = ws1.cell(row=row_idx, column=col_idx, value=value)
                cell.font = body_font
                cell.border = body_border
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                if row_fill:
                    cell.fill = row_fill

        # Column widths
        ws1.column_dimensions["A"].width = 6
        ws1.column_dimensions["B"].width = 80
        ws1.column_dimensions["C"].width = 12
        ws1.freeze_panes = "A4"

        # ── Sheet 2: Word Frequency ──
        ws2 = wb.create_sheet("Word Frequency")
        from collections import Counter
        words = [w.lower().strip(".,!?;:\"'()[]") for ln in lines for w in ln.split() if w.strip()]
        freq  = Counter(words).most_common(50)

        ws2["A1"].value = "Word"
        ws2["B1"].value = "Count"
        for cell in [ws2["A1"], ws2["B1"]]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

        for row_idx, (word, count) in enumerate(freq, start=2):
            ws2.cell(row=row_idx, column=1, value=word).font = body_font
            ws2.cell(row=row_idx, column=2, value=count).font = body_font

        ws2.column_dimensions["A"].width = 30
        ws2.column_dimensions["B"].width = 12

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue(), None

    except Exception as exc:
        logger.exception("XLSX export failed: %s", exc)
        return None, f"XLSX export failed: {exc}"


# ── Notebook-style PNG export ─────────────────────────────────────────────────

def export_notebook_png(text: str) -> ExportResult:
    """
    Renders the extracted text as a high-quality notebook-style image.
    Uses DejaVuSans for clean Unicode rendering.
    Returns PNG bytes.
    """
    try:
        font_path = _font_path_str()
        margin     = 60
        line_height = 28
        font_size   = 16
        title_size  = 20
        bg_color    = (252, 252, 248)
        line_color  = (210, 215, 230)
        text_color  = (30,  35,  50)
        title_color = (80,  70, 200)
        rule_color  = (200, 200, 220)

        try:
            body_font  = ImageFont.truetype(font_path, font_size)
            title_font = ImageFont.truetype(font_path, title_size)
        except Exception:
            body_font  = ImageFont.load_default()
            title_font = ImageFont.load_default()

        # Wrap text to fit within 900px content width
        content_width = 900
        wrapped_lines: list[str] = []
        for raw_line in text.split("\n"):
            if raw_line.strip() == "":
                wrapped_lines.append("")
            else:
                wrapped_lines.extend(textwrap.wrap(raw_line, width=100) or [""])

        # Calculate canvas height
        header_height = 80
        content_height = len(wrapped_lines) * line_height + margin * 2
        total_height = header_height + content_height + margin

        canvas_width  = content_width + margin * 2
        canvas_height = max(total_height, 400)

        img  = Image.new("RGB", (canvas_width, canvas_height), color=bg_color)
        draw = ImageDraw.Draw(img)

        # Subtle ruled lines background
        for y in range(header_height + margin, canvas_height - margin, line_height):
            draw.line([(margin // 2, y), (canvas_width - margin // 2, y)],
                      fill=line_color, width=1)

        # Left margin red line (classic notebook style)
        draw.line([(margin - 10, 0), (margin - 10, canvas_height)],
                  fill=(220, 100, 100), width=2)

        # Title
        draw.text((margin, 20), "Pic2Docs — Extracted Text",
                  font=title_font, fill=title_color)
        draw.text((margin, 46), f"Generated: {_timestamp()}",
                  font=body_font, fill=(140, 140, 160))
        draw.line([(margin // 2, header_height), (canvas_width - margin // 2, header_height)],
                  fill=rule_color, width=1)

        # Body text
        y_offset = header_height + margin
        for line in wrapped_lines:
            if line.strip():
                draw.text((margin, y_offset), line, font=body_font, fill=text_color)
            y_offset += line_height

        buf = io.BytesIO()
        img.save(buf, format="PNG", optimize=True)
        return buf.getvalue(), None

    except Exception as exc:
        logger.exception("PNG export failed: %s", exc)
        return None, f"Notebook PNG export failed: {exc}"
